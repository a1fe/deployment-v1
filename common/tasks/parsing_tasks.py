"""
Parsing tasks for extracting text from resume and job description files
"""

import os
import requests
import re
from datetime import datetime
from typing import Dict, List, Any, Optional
from celery.utils.log import get_task_logger
from dotenv import load_dotenv
from sqlalchemy.orm import Session
from sqlalchemy import text

from common.celery_app.celery_app import celery_app
from common.database.config import database
from common.models.candidates import Submission
from common.models.companies import Job

# Загружаем переменные окружения
load_dotenv()

logger = get_task_logger(__name__)


@celery_app.task(
    bind=True,
    name='common.tasks.parsing_tasks.parse_resume_text',
    soft_time_limit=1800,
    time_limit=2100,
    max_retries=2
)
def parse_resume_text(self, previous_results=None) -> Dict[str, Any]:
    """
    Task 2A: Парсинг текстов резюме из файлов PDF/DOC/DOCX
    
    Обрабатывает все submissions где есть resume_url но нет resume_raw_text
    
    Returns:
        Dict с результатами: количество обработанных файлов, статус
    """
    logger.info("📄 Запуск парсинга текстов резюме")
    
    # Логируем результаты предыдущего этапа
    if previous_results:
        logger.info(f"📥 Получены результаты предыдущего этапа: {previous_results}")
    
    try:
        db = database.get_session()
        processed_count = 0
        error_count = 0
        
        try:
            # Находим submissions с URL файла, но без текста
            submissions_to_process = db.query(Submission).filter(
                Submission.resume_url.isnot(None),
                Submission.resume_url != '',
                Submission.resume_raw_text.is_(None)
            ).all()
            
            logger.info(f"📋 Найдено {len(submissions_to_process)} резюме для парсинга")
            
            for submission in submissions_to_process:
                try:
                    logger.info(f"🔄 Обработка файла резюме: {submission.submission_id}")
                    
                    # Парсим текст из файла
                    extracted_text = _extract_text_from_url(str(submission.resume_url))
                    
                    if extracted_text:
                        # Очищаем текст
                        cleaned_text = _clean_extracted_text(extracted_text)
                        
                        # Сохраняем в БД немедленно после парсинга
                        setattr(submission, 'resume_raw_text', cleaned_text)
                        db.commit()  # Сохраняем сразу для асинхронности
                        processed_count += 1
                        
                        logger.info(f"✅ Текст извлечен и сохранен: {len(cleaned_text)} символов")
                    else:
                        logger.warning(f"⚠️ Не удалось извлечь текст из {submission.resume_url}")
                        error_count += 1
                        
                except Exception as e:
                    logger.error(f"❌ Ошибка обработки резюме {submission.submission_id}: {e}")
                    error_count += 1
                    # Откатываем транзакцию для этого файла
                    db.rollback()
                    continue
            
            # Финальная проверка состояния транзакции
            if db.in_transaction():
                db.commit()
            
            logger.info(f"✅ Парсинг резюме завершен: {processed_count} обработано, {error_count} ошибок")
            
            return {
                'status': 'completed',
                'processed_files': processed_count,
                'error_files': error_count,
                'total_found': len(submissions_to_process)
            }
            
        finally:
            db.close()
            
    except Exception as e:
        logger.error(f"❌ Критическая ошибка парсинга резюме: {e}")
        return {
            'status': 'error',
            'error': str(e),
            'processed_files': 0,
            'error_files': 0
        }


@celery_app.task(
    bind=True,
    name='common.tasks.parsing_tasks.parse_job_text',
    soft_time_limit=1800,
    time_limit=2100,
    max_retries=2
)
def parse_job_text(self, previous_results=None) -> Dict[str, Any]:
    """
    Task 2B: Парсинг текстов описаний вакансий из файлов
    
    Обрабатывает все jobs с job_description_url, но без job_description_raw_text
    
    Returns:
        Dict с результатами: количество обработанных файлов, статус
    """
    logger.info("📄 Запуск парсинга текстов вакансий")
    
    # Логируем результаты предыдущего этапа
    if previous_results:
        logger.info(f"📥 Получены результаты предыдущего этапа: {previous_results}")
    
    try:
        db = database.get_session()
        processed_count = 0
        error_count = 0
        
        try:
            # Находим jobs с URL файла, но без текста
            jobs_to_process = db.query(Job).filter(
                Job.job_description_url.isnot(None),
                Job.job_description_url != '',
                Job.job_description_raw_text.is_(None)
            ).all()
            
            logger.info(f"📋 Найдено {len(jobs_to_process)} вакансий для парсинга")
            
            for job in jobs_to_process:
                try:
                    logger.info(f"🔄 Обработка файла вакансии: {job.job_id}")
                    
                    # Парсим текст из файла
                    extracted_text = _extract_text_from_url(str(job.job_description_url))
                    
                    if extracted_text:
                        # Очищаем текст
                        cleaned_text = _clean_extracted_text(extracted_text)
                        
                        # Сохраняем в БД немедленно после парсинга
                        setattr(job, 'job_description_raw_text', cleaned_text)
                        db.commit()  # Сохраняем сразу для асинхронности
                        processed_count += 1
                        
                        logger.info(f"✅ Текст извлечен и сохранен: {len(cleaned_text)} символов")
                    else:
                        logger.warning(f"⚠️ Не удалось извлечь текст из {job.job_description_url}")
                        error_count += 1
                        
                except Exception as e:
                    logger.error(f"❌ Ошибка обработки вакансии {job.job_id}: {e}")
                    error_count += 1
                    # Откатываем транзакцию для этого файла
                    db.rollback()
                    continue
            
            # Финальная проверка состояния транзакции
            if db.in_transaction():
                db.commit()
            
            logger.info(f"✅ Парсинг вакансий завершен: {processed_count} обработано, {error_count} ошибок")
            
            return {
                'status': 'completed',
                'processed_files': processed_count,
                'error_files': error_count,
                'total_found': len(jobs_to_process)
            }
            
        finally:
            db.close()
            
    except Exception as e:
        logger.error(f"❌ Критическая ошибка парсинга вакансий: {e}")
        return {
            'status': 'error',
            'error': str(e),
            'processed_files': 0,
            'error_files': 0
        }


def _extract_text_from_url(file_url: str) -> Optional[str]:
    """
    Извлекает текст из файла по URL
    
    Args:
        file_url: URL файла для парсинга
        
    Returns:
        Извлеченный текст или None при ошибке
    """
    if not file_url or not file_url.startswith(('http://', 'https://')):
        logger.warning(f"⚠️ Некорректный URL: {file_url}")
        return None
    
    try:
        logger.info(f"📥 Загрузка файла: {file_url}")
        
        # Загружаем файл
        response = requests.get(file_url, timeout=30)
        response.raise_for_status()
        
        # Определяем тип файла по URL или Content-Type
        content_type = response.headers.get('content-type', '').lower()
        file_extension = _get_file_extension(file_url, content_type)
        
        logger.info(f"📋 Тип файла: {file_extension}, Размер: {len(response.content)} байт")
        
        # Парсим в зависимости от типа файла
        if file_extension == 'pdf':
            return _extract_pdf_text(response.content)
        elif file_extension in ['doc', 'docx']:
            return _extract_docx_text(response.content)
        else:
            logger.warning(f"⚠️ Неподдерживаемый тип файла: {file_extension}")
            return None
            
    except requests.RequestException as e:
        logger.error(f"❌ Ошибка загрузки файла {file_url}: {e}")
        return None
    except Exception as e:
        logger.error(f"❌ Ошибка обработки файла {file_url}: {e}")
        return None


def _get_file_extension(url: str, content_type: str) -> str:
    """Определяет расширение файла по URL или Content-Type"""
    # Сначала пробуем по URL
    if url.lower().endswith('.pdf'):
        return 'pdf'
    elif url.lower().endswith(('.doc', '.docx')):
        return 'docx'
    
    # Затем по Content-Type
    if 'pdf' in content_type:
        return 'pdf'
    elif 'word' in content_type or 'officedocument' in content_type:
        return 'docx'
    elif 'msword' in content_type:
        return 'doc'
    
    # По умолчанию пробуем PDF
    return 'pdf'


def _extract_pdf_text(pdf_content: bytes) -> Optional[str]:
    """
    Извлекает текст из PDF файла с fallback методами
    
    Args:
        pdf_content: Содержимое PDF файла в байтах
        
    Returns:
        Извлеченный текст или None
    """
    # Пробуем PyMuPDF сначала
    result = _extract_pdf_with_fitz(pdf_content)
    if result:
        return result
    
    # Fallback на PyPDF2
    logger.info("🔄 Переход на PyPDF2 для извлечения текста")
    return _extract_pdf_with_pypdf2(pdf_content)


def _extract_pdf_with_fitz(pdf_content: bytes) -> Optional[str]:
    """Извлечение текста с помощью PyMuPDF"""
    doc = None
    try:
        import fitz  # PyMuPDF
        
        # Открываем PDF из памяти
        doc = fitz.open(stream=pdf_content, filetype="pdf")
        text_parts = []
        page_count = doc.page_count
        
        # Извлекаем текст со всех страниц
        for page_num in range(page_count):
            try:
                page = doc[page_num]
                # Используем стандартный API PyMuPDF для извлечения текста
                text = page.get_text()  # type: ignore
                if text and text.strip():
                    text_parts.append(text.strip())
            except Exception as page_error:
                logger.warning(f"⚠️ Ошибка извлечения текста со страницы {page_num + 1}: {page_error}")
                continue
        
        if text_parts:
            full_text = '\n\n'.join(text_parts)
            logger.info(f"📄 PDF обработан PyMuPDF: {page_count} страниц, {len(full_text)} символов")
            return full_text
        else:
            logger.warning("⚠️ PyMuPDF: PDF не содержит извлекаемого текста")
            return None
            
    except ImportError:
        logger.warning("⚠️ PyMuPDF не доступен, используем fallback")
        return None
    except Exception as e:
        logger.warning(f"⚠️ Ошибка PyMuPDF: {e}, используем fallback")
        return None
    finally:
        # Гарантированно закрываем документ
        if doc is not None:
            try:
                doc.close()
            except Exception as close_error:
                logger.warning(f"⚠️ Ошибка закрытия PDF: {close_error}")


def _extract_pdf_with_pypdf2(pdf_content: bytes) -> Optional[str]:
    """Fallback извлечение текста с помощью PyPDF2"""
    try:
        import PyPDF2
        from io import BytesIO
        
        pdf_stream = BytesIO(pdf_content)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
            except Exception as page_error:
                logger.warning(f"⚠️ PyPDF2: Ошибка извлечения текста со страницы {page_num + 1}: {page_error}")
                continue
        
        if text_parts:
            full_text = '\n\n'.join(text_parts)
            logger.info(f"📄 PDF обработан PyPDF2: {len(pdf_reader.pages)} страниц, {len(full_text)} символов")
            return full_text
        else:
            logger.warning("⚠️ PyPDF2: PDF не содержит извлекаемого текста")
            return None
            
    except ImportError:
        logger.error("❌ PyPDF2 не установлен. Установите: pip install PyPDF2")
        return None
    except Exception as e:
        logger.error(f"❌ Ошибка PyPDF2: {e}")
        return None


def _extract_docx_text(docx_content: bytes) -> Optional[str]:
    """
    Извлекает текст из DOCX/DOC файла с поддержкой разных форматов
    
    Args:
        docx_content: Содержимое DOCX/DOC файла в байтах
        
    Returns:
        Извлеченный текст или None
    """
    # Сначала пробуем стандартный DOCX формат
    result = _extract_modern_docx(docx_content)
    if result:
        return result
    
    # Fallback для старых DOC файлов или поврежденных DOCX
    logger.info("🔄 Пробуем альтернативные методы извлечения текста из DOC")
    return _extract_legacy_doc_text(docx_content)


def _extract_modern_docx(docx_content: bytes) -> Optional[str]:
    """Извлечение текста из современного DOCX формата"""
    try:
        import docx
        from io import BytesIO
        
        # Открываем DOCX из памяти
        doc = docx.Document(BytesIO(docx_content))
        text_parts = []
        
        # Извлекаем текст из всех параграфов
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)
        
        if text_parts:
            full_text = '\n'.join(text_parts)
            logger.info(f"📄 DOCX обработан: {len(doc.paragraphs)} параграфов, {len(full_text)} символов")
            return full_text
        else:
            logger.warning("⚠️ DOCX не содержит текста")
            return None
            
    except ImportError:
        logger.error("❌ python-docx не установлен. Установите: pip install python-docx")
        return None
    except Exception as e:
        # Обрабатываем любые ошибки, включая проблемы с форматом файла
        if "not a zip file" in str(e).lower() or "bad zipfile" in str(e).lower():
            logger.warning("⚠️ Файл не является современным DOCX документом, пробуем legacy методы")
        else:
            logger.warning(f"⚠️ Ошибка обработки DOCX: {e}, пробуем альтернативные методы")
        return None


def _extract_legacy_doc_text(doc_content: bytes) -> Optional[str]:
    """Fallback извлечение текста из старых DOC файлов или поврежденных документов"""
    try:
        # Попытка использовать antiword для DOC файлов
        text_result = _extract_with_antiword(doc_content)
        if text_result:
            return text_result
        
        # Попытка использовать olefile для чтения OLE структуры DOC
        text_result = _extract_with_olefile(doc_content)
        if text_result:
            return text_result
        
        # Попытка простого текстового поиска в бинарных данных
        # (работает для некоторых старых DOC файлов)
        text_content = _extract_text_from_binary(doc_content)
        if text_content:
            logger.info(f"📄 DOC обработан бинарным методом: {len(text_content)} символов")
            return text_content
        
        logger.error("❌ Не удалось извлечь текст ни одним из доступных методов")
        return None
        
    except Exception as e:
        logger.error(f"❌ Критическая ошибка при обработке legacy DOC: {e}")
        return None


def _extract_with_antiword(doc_content: bytes) -> Optional[str]:
    """Извлечение текста с помощью antiword библиотеки"""
    try:
        import tempfile
        import subprocess
        import os
        
        # antiword работает как командная утилита, создаем временный файл
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
            temp_file.write(doc_content)
            temp_file.flush()
            
            try:
                # Пробуем использовать antiword через subprocess
                try:
                    result = subprocess.run(
                        ['antiword', temp_file.name],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        text = result.stdout.strip()
                        logger.info(f"📄 DOC обработан antiword CLI: {len(text)} символов")
                        return text
                except (subprocess.TimeoutExpired, FileNotFoundError):
                    logger.warning("⚠️ antiword CLI недоступен")
                    
            finally:
                # Удаляем временный файл
                try:
                    os.unlink(temp_file.name)
                except:
                    pass
                    
    except Exception as e:
        logger.warning(f"⚠️ antiword не смог обработать файл: {e}")
        return None


def _extract_with_olefile(doc_content: bytes) -> Optional[str]:
    """Извлечение текста с помощью olefile для чтения OLE структуры"""
    try:
        import olefile
        from io import BytesIO
        
        # Проверяем, является ли файл OLE документом
        if not olefile.isOleFile(BytesIO(doc_content)):
            logger.warning("⚠️ Файл не является OLE документом")
            return None
        
        ole = olefile.OleFileIO(BytesIO(doc_content))
        
        try:
            # Получаем список всех доступных потоков
            all_streams = ole.listdir()
            logger.info(f"🔍 Найденные потоки в DOC файле: {all_streams[:10]}...")  # Показываем первые 10
            
            found_text = []
            
            # Приоритетные потоки для поиска текста
            priority_streams = ['WordDocument', '1Table', 'Data', 'CompObj']
            
            # Проверяем приоритетные потоки
            for stream_name in priority_streams:
                if ole.exists(stream_name):
                    try:
                        stream_data = ole.openstream(stream_name).read()
                        
                        if stream_data and len(stream_data) > 100:
                            text_parts = _extract_readable_text_from_bytes(stream_data)
                            if text_parts:
                                found_text.extend(text_parts)
                                logger.info(f"� Найден текст в потоке {stream_name}: {len(text_parts)} частей")
                                
                    except Exception as stream_error:
                        logger.warning(f"⚠️ Ошибка чтения потока {stream_name}: {stream_error}")
                        continue
            
            # Если ничего не нашли в приоритетных потоках, попробуем другие
            if not found_text:
                for stream_info in all_streams:
                    try:
                        stream_name = stream_info[0] if isinstance(stream_info, list) else stream_info
                        
                        # Пропускаем системные потоки
                        if stream_name.startswith(('__', '\x01', '\x03', '\x05')):
                            continue
                            
                        if ole.exists(stream_name):
                            stream_data = ole.openstream(stream_name).read()
                            
                            if stream_data and len(stream_data) > 100:
                                text_parts = _extract_readable_text_from_bytes(stream_data)
                                if text_parts:
                                    found_text.extend(text_parts)
                                    logger.info(f"📄 Найден текст в потоке {stream_name}: {len(text_parts)} частей")
                                    # Если нашли достаточно текста, можем остановиться
                                    if len(' '.join(found_text)) > 500:
                                        break
                                        
                    except Exception:
                        continue
            
            if found_text:
                text = ' '.join(found_text)
                text = re.sub(r'\s+', ' ', text).strip()
                if len(text) > 50:
                    logger.info(f"📄 DOC обработан olefile: {len(text)} символов")
                    return text
                        
        finally:
            ole.close()
            
    except ImportError:
        logger.warning("⚠️ olefile не доступен")
        return None
    except Exception as e:
        logger.warning(f"⚠️ olefile не смог обработать файл: {e}")
        return None


def _extract_readable_text_from_bytes(data: bytes) -> List[str]:
    """Извлекает читаемый текст из байтовых данных"""
    text_parts = []
    
    # Пробуем разные кодировки
    encodings = ['utf-16le', 'utf-16be', 'utf-8', 'latin-1', 'cp1252', 'cp1251']
    
    for encoding in encodings:
        try:
            decoded = data.decode(encoding, errors='ignore')
            # Ищем читаемые текстовые части (минимум 5 символов)
            import re
            readable_parts = re.findall(r'[A-Za-z][A-Za-z0-9\s\.,;:!?\-()]{4,}', decoded)
            
            # Фильтруем осмысленные части
            valid_parts = []
            for part in readable_parts:
                # Проверяем соотношение букв к общему количеству символов
                letters = len(re.findall(r'[A-Za-z]', part))
                if letters > len(part) * 0.3:  # Минимум 30% букв
                    valid_parts.append(part.strip())
            
            if valid_parts:
                text_parts.extend(valid_parts)
                
        except Exception:
            continue
    
    # Удаляем дубликаты и очень короткие части
    unique_parts = []
    for part in text_parts:
        if len(part) > 10 and part not in unique_parts:
            unique_parts.append(part)
    
    return unique_parts


def _extract_text_from_binary(binary_content: bytes) -> Optional[str]:
    """Простое извлечение текста из бинарных данных DOC файла"""
    try:
        # Пробуем разные кодировки
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                # Декодируем с игнорированием ошибок
                raw_text = binary_content.decode(encoding, errors='ignore')
                
                # Извлекаем читаемый текст (минимум 3 символа подряд)
                import re
                text_parts = re.findall(r'[A-Za-z0-9\s\.,;:!?\-()]{3,}', raw_text)
                
                if text_parts:
                    # Объединяем найденные части
                    extracted_text = ' '.join(text_parts)
                    
                    # Очищаем от лишних пробелов
                    extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                    
                    # Проверяем, что получился осмысленный текст
                    if len(extracted_text) > 50 and len(extracted_text.split()) > 10:
                        logger.info(f"📄 Извлечен текст кодировкой {encoding}: {len(extracted_text)} символов")
                        return extracted_text
                        
            except Exception as e:
                continue
        
        logger.warning("⚠️ Не удалось извлечь читаемый текст из бинарных данных")
        return None
        
    except Exception as e:
        logger.error(f"❌ Ошибка бинарного извлечения текста: {e}")
        return None


def _clean_extracted_text(raw_text: str) -> str:
    """
    Очищает извлеченный текст от лишних элементов
    
    Args:
        raw_text: Сырой извлеченный текст
        
    Returns:
        Очищенный текст
    """
    if not raw_text:
        return ""
    
    # Нормализация переносов строк
    text = re.sub(r'\r\n|\r', '\n', raw_text)
    
    # Удаление лишних пробелов и табуляций
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Удаление множественных переносов строк (больше 2 подряд)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Удаление пробелов в начале и конце строк
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Удаление дублированных строк (простое удаление идентичных соседних строк)
    lines = text.split('\n')
    cleaned_lines = []
    previous_line = None
    
    for line in lines:
        if line != previous_line or len(line.strip()) > 50:  # Оставляем длинные строки даже если дублированы
            cleaned_lines.append(line)
        previous_line = line
    
    text = '\n'.join(cleaned_lines)
    
    # Финальная очистка
    text = text.strip()
    
    # Удаляем очень короткие тексты (менее 10 символов)
    if len(text) < 10:
        logger.warning("⚠️ Извлеченный текст слишком короткий")
        return ""
    
    logger.info(f"🧹 Текст очищен: {len(raw_text)} → {len(text)} символов")
    return text
